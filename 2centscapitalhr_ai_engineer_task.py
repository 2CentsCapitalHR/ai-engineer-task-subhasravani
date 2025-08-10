

# Install necessary libraries
!pip install python-docx==0.8.11 sentence-transformers faiss-cpu PyPDF2 streamlit --quiet

# Imports and basic helpers
import os, io, json, re
from collections import defaultdict
from docx import Document
from docx.shared import RGBColor
from google.colab import files
from sentence_transformers import SentenceTransformer
import numpy as np
import faiss
import PyPDF2
import streamlit as st
import datetime
import subprocess
import threading
import time
import socket
# Optional for LLM calls
import requests

# Simple utility to save JSON and download
def save_json(obj, path="report.json"):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(obj, f, indent=2, ensure_ascii=False)
    return path



GEMINI_API_KEY = "AIzaSyDPlJn5jQx0p8Svdv4KtkG2bHV0CJI-jXA"

import google.generativeai as genai
genai.configure(api_key=GEMINI_API_KEY)

print("API key setup complete.")

"""### Document Parsing and Classification
Functions to extract text from .docx and classify documents.
"""

def get_docx_text(docx_path):
    """
    Extracts text from a .docx file.
    """
    document = Document(docx_path)
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)

def classify_document(text):
    """
    Basic document classification based on keywords.
    
    """
    text_lower = text.lower()
    if "articles of association" in text_lower or "memorandum of association" in text_lower:
        return "AoA/MoA"
    # Add more classification rules as needed
    return "Other"

"""### ADGM Checklist and RAG Placeholder Functions
Define the checklist and placeholder functions for RAG-based comparison.
"""

# This is a placeholder checklist.
adgm_checklist = [
    "Company name includes 'Limited' or 'Ltd.' (if applicable)",
    "Registered office address is within ADGM",
    "Objects clause is clearly defined",
    "Share capital structure is compliant with ADGM regulations",
    "Details of directors are provided",
    "Details of shareholders are provided",
    "Articles of Association are present and comply with ADGM rules",
    # Add more checklist items based on ADGM requirements
]

def retrieve_adgm_context(query, index, documents, model, k=3):
    """
    Retrieves relevant ADGM context based on a query using RAG.
    (Placeholder function - needs implementation with actual RAG logic)
    """
    # Placeholder RAG: Simple keyword matching (replace with actual RAG using index and model)
    if index is None or documents is None or model is None:
         # Fallback to simple keyword matching if RAG components not provided
        relevant_docs = []
        for doc in documents: 
            if query.lower() in doc.lower():
                relevant_docs.append(doc)
        return "\n".join(relevant_docs)
    else:
        # Actual RAG implementation using index and model
        query_embedding = model.encode(query)
        query_embedding = np.array([query_embedding])
        distances, ids = index.search(query_embedding, k)
        relevant_docs = [documents[id] for id in ids[0]]
        return "\n".join(relevant_docs)


def compare_with_checklist(document_text, checklist, rag_index=None, rag_documents=None, rag_model=None):
    """
    Compares the document text against the ADGM checklist and detects red flags.
    Uses RAG to get relevant context for verification.
    """
    red_flags = []
    checklist_results = {}

    for item in checklist:
        # Use RAG to get relevant ADGM rules for the checklist item
        adgm_rules_context = ""
        if rag_index and rag_documents and rag_model:
             adgm_rules_context = retrieve_adgm_context(item, rag_index, rag_documents, rag_model) # Pass RAG components

      
        prompt = f"""
        Review the following document text and the provided ADGM rules context.
        Determine if the document text satisfies the checklist item based *primarily* on the ADGM rules context.

        Document Text:
        ---
        {document_text[:2000]} 
        ---

        Checklist Item: "{item}"

        ADGM Rules Context:
        ---
        {adgm_rules_context}
        ---

        Based on the Document Text and explicitly referencing the ADGM Rules Context, does the document satisfy the checklist item?
        Provide a concise answer: "Yes" if compliant according to the ADGM Rules Context, or "No" if not compliant or if the ADGM Rules Context indicates a requirement not met in the Document Text.
        If the answer is "No", state the specific red flag or reason for non-compliance, citing the relevant part of the ADGM Rules Context if possible.

        Answer Format:
        Compliance: [Yes/No]
        Red Flag (if No): [Specific reason and relevant ADGM rule if applicable]
        """
       

        # Placeholder result - Replace with actual LLM call and parsing
        result_text = "Compliance: Needs manual verification.\nRed Flag (if No): Unable to verify with placeholder LLM."
        red_flag_found = "Possible red flag: Needs detailed review against ADGM rules." # Keep placeholder for now

        # Basic parsing of placeholder result (replace with proper LLM response parsing)
        compliance_status = "Needs manual verification."
        flag_detail = "Unable to verify with placeholder LLM."
        for line in result_text.split('\n'):
            if line.startswith("Compliance:"):
                compliance_status = line.split(":")[1].strip()
            elif line.startswith("Red Flag (if No):"):
                flag_detail = line.split(":")[1].strip()

        checklist_results[item] = compliance_status
        if compliance_status.lower() == "no" or "red flag" in flag_detail.lower():
             red_flags.append(f"{item}: {flag_detail}") # Customize red flag message based on LLM output

    return checklist_results, red_flags

"""### Load and Process ADGM Reference Documents
Load the PDF, chunk it, create embeddings, and build the FAISS index for RAG. **Ensure "Data Sources.pdf" is in `/content/`**.
"""

# Define the path to your ADGM reference PDF document
adgm_reference_pdf_path = "/content/Data Sources.pdf"
adgm_reference_text = ""

try:
    if not os.path.exists(adgm_reference_pdf_path):
        raise FileNotFoundError(f"ADGM reference PDF not found at {adgm_reference_pdf_path}")

    with open(adgm_reference_pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            adgm_reference_text += page.extract_text()
    print("Successfully extracted text from ADGM reference PDF.")

    # Chunk the text
    chunk_size = 1000
    adgm_chunks = []
    for i in range(0, len(adgm_reference_text), chunk_size):
        adgm_chunks.append(adgm_reference_text[i:i + chunk_size])
    print(f"Split text into {len(adgm_chunks)} chunks.")

    # Create embeddings
    model = SentenceTransformer('all-MiniLM-L6-v2')
    adgm_embeddings = model.encode(adgm_chunks)
    print(f"Shape of the generated embeddings: {adgm_embeddings.shape}")

    # Build FAISS index
    embedding_dimensionality = adgm_embeddings.shape[1]
    index = faiss.IndexFlatL2(embedding_dimensionality)
    index = faiss.IndexIDMap(index)
    index.add_with_ids(adgm_embeddings, np.array(range(len(adgm_chunks))))
    print(f"Added {index.ntotal} vectors to the FAISS index.")

except FileNotFoundError as fnf_error:
    print(f"Error: {fnf_error}. Please upload 'Data Sources.pdf' to /content/")
    # Initialize placeholders if file not found, to prevent errors in later cells
    adgm_chunks = []
    adgm_embeddings = None
    index = None
    model = None
except Exception as e:
    print(f"An error occurred during RAG setup: {e}")
    # Initialize placeholders in case of other errors
    adgm_chunks = []
    adgm_embeddings = None
    index = None
    model = None

"""### Output Generation Functions
Functions to add comments (placeholder) and generate the summary.
"""

def add_comment_to_docx(document, text_to_comment, comment_text):
    """
    Adds an inline comment to a specific text within a .docx document.
    This is a simplified example and may require more sophisticated logic
    to find and precisely comment on specific text spans.
    """
    for paragraph in document.paragraphs:
        if text_to_comment in paragraph.text:
            # This is a basic placeholder. Real implementation needs careful handling
            # of runs and text matching to insert comments correctly.
            # docx library's commenting feature is complex and might require xml manipulation.
            print(f"Placeholder: Would add comment '{comment_text}' near '{text_to_comment}'")
            # Example of highlighting (simpler than commenting):
            # for run in paragraph.runs:
            #     if text_to_comment in run.text:
            #         run.font.highlight_color = 6 # Yellow highlight


def generate_summary(filename, document_type, checklist_results, red_flags):
    """
    Generates a structured summary of the document review.
    """
    summary = {
        "filename": filename,
        "document_type": document_type,
        "checklist_results": checklist_results,
        "red_flags": red_flags,
        "review_timestamp": datetime.datetime.now().isoformat()
    }
    return summary

"""### Streamlit Application Code
This code will be written to `app.py` to create the web UI.
"""

app_code = """
import streamlit as st
import os
from docx import Document
import datetime
import json
from sentence_transformers import SentenceTransformer
import numpy as np
import faiss
import PyPDF2 # Import PyPDF2 here for PDF loading in Streamlit

# --- Helper Functions (Copy/Paste from previous cells or import) ---

def get_docx_text(docx_path):
    \"\"\"
    Extracts text from a .docx file.
    \"\"\"
    document = Document(docx_path)
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text)
    return "\\n".join(text)

def classify_document(text):
    \"\"\"
    Basic document classification based on keywords.
    Can be extended with more sophisticated methods (e.g., ML models).
    \"\"\"
    text_lower = text.lower()
    if "articles of association" in text_lower or "memorandum of association" in text_lower:
        return "AoA/MoA"
    # Add more classification rules as needed
    return "Other"

# This is a placeholder checklist. You will need to define your actual ADGM checklist.
adgm_checklist = [
    "Company name includes 'Limited' or 'Ltd.' (if applicable)",
    "Registered office address is within ADGM",
    "Objects clause is clearly defined",
    "Share capital structure is compliant with ADGM regulations",
    "Details of directors are provided",
    "Details of shareholders are provided",
    "Articles of Association are present and comply with ADGM rules",
    # Add more checklist items based on ADGM requirements
]

# RAG function within Streamlit app
def retrieve_adgm_context_streamlit(query, index, documents, model, k=3):
    \"\"\"
    Retrieves relevant ADGM context based on a query using RAG.
    (Uses RAG components passed or assumed available)
    \"\"\"
    if index is None or documents is None or model is None:
        # This fallback should ideally not be reached if RAG components are loaded in Streamlit
        return "ADGM reference data not loaded or RAG setup failed."
    try:
        query_embedding = model.encode(query)
        query_embedding = np.array([query_embedding]) # Ensure it's 2D
        distances, ids = index.search(query_embedding, k)
        relevant_docs = [documents[id] for id in ids[0]]
        return "\\n".join(relevant_docs)
    except Exception as e:
        return f"Error during RAG retrieval: {e}"


def compare_with_checklist_streamlit(document_text, checklist, rag_index=None, rag_documents=None, rag_model=None):
    \"\"\"
    Compares the document text against the ADGM checklist and detects red flags.
    Uses RAG to get relevant context for verification.
    \"\"\"
    red_flags = []
    checklist_results = {}

    for item in checklist:
        # Use RAG to get relevant ADGM rules for the checklist item
        # Pass the RAG components loaded within the Streamlit app
        adgm_rules_context = retrieve_adgm_context_streamlit(item, rag_index, rag_documents, rag_model)

        # Use LLM to check if the document meets the checklist item based on context
        # This is a simplified example using a placeholder. Replace with actual LLM integration.
        prompt = f\"\"\"
        Review the following document text and the provided ADGM rules context.
        Determine if the document text satisfies the checklist item based *primarily* on the ADGM rules context.

        Document Text:
        ---
        {document_text[:2000]} # Use a reasonable chunk of text
        ---

        Checklist Item: "{item}"

        ADGM Rules Context:
        ---
        {adgm_rules_context}
        ---

        Based on the Document Text and explicitly referencing the ADGM Rules Context, does the document satisfy the checklist item?
        Provide a concise answer: "Yes" if compliant according to the ADGM Rules Context, or "No" if not compliant or if the ADGM Rules Context indicates a requirement not met in the Document Text.
        If the answer is "No", state the specific red flag or reason for non-compliance, citing the relevant part of the ADGM Rules Context if possible.

        Answer Format:
        Compliance: [Yes/No]
        Red Flag (if No): [Specific reason and relevant ADGM rule if applicable]
        \"\"\"
        # Replace with your actual LLM call (e.g., using genai.GenerativeModel)
        # Ensure genai is imported and configured in the Streamlit app or accessible.
        # try:
        #     import google.generativeai as genai
        #     # Configure API key if not done globally or passed
        #     # genai.configure(api_key="YOUR_STREAMLIT_API_KEY") # Consider secure handling
        #     llm_model = genai.GenerativeModel('gemini-1.5-flash') # Or your preferred model
        #     response = llm_model.generate_content(prompt)
        #     result_text = response.text.strip()
        # except Exception as e:
        #     result_text = f"LLM call failed: {e}" # Handle LLM errors

        # Placeholder result - Replace with actual LLM call and parsing
        result_text = "Compliance: Needs manual verification.\\nRed Flag (if No): Unable to verify with placeholder LLM." # Escaped newline
        red_flag_found = "Possible red flag: Needs detailed review against ADGM rules." # Keep placeholder for now

        # Basic parsing of placeholder result (replace with proper LLM response parsing)
        compliance_status = "Needs manual verification."
        flag_detail = "Unable to verify with placeholder LLM."
        for line in result_text.split('\\n'): # Use \\n for newline in string literal
            if line.startswith("Compliance:"):
                compliance_status = line.split(":")[1].strip()
            elif line.startswith("Red Flag (if No):"):
                flag_detail = line.split(":")[1].strip()

        checklist_results[item] = compliance_status
        if compliance_status.lower() == "no" or "red flag" in flag_detail.lower():
             red_flags.append(f"{item}: {flag_detail}") # Customize red flag message based on LLM output

    return checklist_results, red_flags


def add_comment_to_docx(document, text_to_comment, comment_text):
    \"\"\"
    Adds an inline comment to a specific text within a .docx document.
    This is a simplified example and may require more sophisticated logic
    to find and precisely comment on specific text spans.
    \"\"\"
    for paragraph in document.paragraphs:
        if text_to_comment in paragraph.text:
            # This is a basic placeholder. Real implementation needs careful handling
            # of runs and text matching to insert comments correctly.
            # docx library's commenting feature is complex and might require xml manipulation.
            print(f"Placeholder: Would add comment '{comment_text}' near '{text_to_comment}'")
            # Example of highlighting (simpler than commenting):
            # for run in paragraph.runs:
            #     if text_to_comment in run.text:
            #         run.font.highlight_color = 6 # Yellow highlight


def generate_summary(filename, document_type, checklist_results, red_flags):
    \"\"\"
    Generates a structured summary of the document review.
    \"\"\"
    summary = {
        "filename": filename,
        "document_type": document_type,
        "checklist_results": checklist_results,
        "red_flags": red_flags,
        "review_timestamp": datetime.datetime.now().isoformat()
    }
    return summary

def save_json(obj, path="report.json"):
    \"\"\"
    Simple utility to save JSON and download.
    \"\"\"
    with open(path, "w", encoding="utf-8") as f:
        json.dump(obj, f, indent=2, ensure_ascii=False)
    return path


# --- Streamlit App ---

st.title("ADGM-Compliant Corporate Agent")

# --- RAG Setup within Streamlit App ---
# Define the path to your ADGM reference PDF document
adgm_reference_pdf_path = "/content/Data Sources.pdf" # Ensure this path is correct in Colab

@st.cache_resource # Cache the RAG components to avoid reloading on each interaction
def setup_rag(pdf_path):
    adgm_reference_text = ""
    adgm_chunks = []
    adgm_embeddings = None
    index = None
    model = None
    rag_load_status = "Not loaded"

    try:
        if not os.path.exists(pdf_path):
            rag_load_status = f"Error: ADGM reference PDF not found at {pdf_path}. RAG not set up."
            st.error(rag_load_status)
            return None, [], None, None, rag_load_status # Return components and status

        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                adgm_reference_text += page.extract_text()
        st.success("Successfully extracted text from ADGM reference PDF.")
        rag_load_status = "PDF text extracted."

        # Chunk the text
        chunk_size = 1000
        adgm_chunks = []
        for i in range(0, len(adgm_reference_text), chunk_size):
            adgm_chunks.append(adgm_reference_text[i:i + chunk_size])
        st.success(f"Split text into {len(adgm_chunks)} chunks.")
        rag_load_status += f" Split into {len(adgm_chunks)} chunks."


        # Create embeddings
        model = SentenceTransformer('all-MiniLM-L6-v2')
        adgm_embeddings = model.encode(adgm_chunks)
        st.success(f"Shape of the generated embeddings: {adgm_embeddings.shape}")
        rag_load_status += f" Embeddings shape: {adgm_embeddings.shape}."


        # Build FAISS index
        embedding_dimensionality = adgm_embeddings.shape[1]
        index = faiss.IndexFlatL2(embedding_dimensionality)
        index = faiss.IndexIDMap(index)
        index.add_with_ids(adgm_embeddings, np.array(range(len(adgm_chunks))))
        st.success(f"Added {index.ntotal} vectors to the FAISS index.")
        rag_load_status += f" FAISS index built with {index.ntotal} vectors."

        rag_load_status = "RAG setup complete!"
        return index, adgm_chunks, model, adgm_embeddings, rag_load_status # Return all components and final status

    except Exception as e:
        rag_load_status = f"An error occurred during RAG setup in Streamlit: {e}"
        st.error(rag_load_status)
        return None, [], None, None, rag_load_status # Return components and status in case of error

# Call the setup_rag function when the Streamlit app starts
index, adgm_chunks, model, adgm_embeddings, rag_status_message = setup_rag(adgm_reference_pdf_path)

# Display RAG load status in the sidebar
st.sidebar.text(f"RAG Status: {rag_status_message}")

# --- Main App Logic ---

uploaded_file = st.file_uploader("Upload a .docx legal document", type=["docx"])


if uploaded_file is not None:
    # Save the uploaded file temporarily
    file_path = os.path.join("/tmp", uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    st.write(f"Processing file: {uploaded_file.name}")

    # 1. Parse and Classify
    document_text = get_docx_text(file_path)
    document_type = classify_document(document_text)
    st.write(f"Document classified as: {document_type}")

    # 2. Compare with Checklist and Detect Red Flags
    st.write("Comparing document with ADGM checklist...")
    # Pass the RAG components loaded by setup_rag
    checklist_results, red_flags = compare_with_checklist_streamlit(
        document_text, adgm_checklist, index, adgm_chunks, model
    )

    st.subheader("Checklist Results:")
    for item, result in checklist_results.items():
        st.write(f"- {item}: {result}")

    st.subheader("Detected Red Flags:")
    if red_flags:
        for flag in red_flags:
            st.error(f"- {flag}")
    else:
        st.success("No major red flags detected based on the checklist.")

    # 3. Generate Output (Document with Comments and Summary)
    st.write("Generating reviewed document and summary...")

    # Load the document again to add comments (replace with your actual comment logic)
    reviewed_document = Document(file_path)
    # Example of adding a placeholder comment
    # add_comment_to_docx(reviewed_document, "Company name", "Verify company name format with ADGM rules.")

    output_document_path = f"reviewed_{uploaded_file.name}"
    reviewed_document.save(output_document_path)

    review_summary = generate_summary(uploaded_file.name, document_type, checklist_results, red_flags)
    # Save summary in the same directory as the app for easier access via download button
    summary_json_path = f"summary_{uploaded_file.name}.json"
    save_json(review_summary, summary_json_path)


    st.subheader("Outputs:")
    st.download_button(
        label="Download Reviewed Document (.docx)",
        data=open(output_document_path, "rb"),
        file_name=output_document_path,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.download_button(
        label="Download Review Summary (.json)",
        data=open(summary_json_path, "r").read(),
        file_name=summary_json_path,
        mime="application/json"
    )

    st.success("Processing complete!")

"""

with open("app.py", "w") as f:
    f.write(app_code)

print("app.py created successfully in the /content/ directory.")

"""### Run Streamlit Application
Execute this cell to start the Streamlit app and get the public URL.
"""

!streamlit run app.py & npx localtunnel --port 8501

"""## Summary of Task and Challenges

### Task Overview:
The goal is to build an intelligent AI-powered legal assistant called the Corporate Agent for reviewing and validating documentation for business incorporation and compliance within the Abu Dhabi Global Market (ADGM) jurisdiction. Key capabilities include accepting .docx documents, verifying completeness based on ADGM rules, highlighting red flags, inserting contextual comments citing ADGM rules, generating a reviewed downloadable file and a structured JSON/Python report, and using RAG with provided ADGM reference documents for accuracy. The agent should also inform users if required documents are missing and have a UI in Gradio or Streamlit.

### Summary of Work Done:
*   Set up project structure and installed necessary libraries (`python-docx`, `sentence-transformers`, `faiss-cpu`, `PyPDF2`, `streamlit`).
*   Implemented functions for .docx parsing and basic document classification.
*   Defined the ADGM checklist and created placeholder functions for RAG-based comparison and red flag detection.
*   Implemented core RAG steps: loading text from "Data Sources.pdf" (after file availability was resolved), chunking text, creating embeddings, and building a FAISS index.
*   Integrated the RAG retrieval logic into the checklist comparison function.
*   Updated the LLM prompt within the comparison function to utilize the retrieved ADGM context (actual LLM calls are still placeholders).
*   Created functions for generating a JSON summary and a placeholder for adding comments to the .docx file.
*   Developed the Streamlit web application code (`app.py`) for the UI.
*   Modified the Streamlit app to load and build RAG components internally to resolve access issues.
*   Provided the command to run the Streamlit app in Colab via `localtunnel`.

### Challenges Faced:
*   **File Not Found:** Repeated issues with locating and accessing the "Data Sources.pdf" file.
*   **Streamlit `app.py` Location and Syntax:** Problems ensuring `app.py` was created in the correct directory and fixing a `SyntaxError` in the generated code.
*   **RAG Component Access in Streamlit:** The Streamlit process could not directly access variables from the notebook kernel, requiring RAG setup to be moved inside the Streamlit app.
*   **Placeholder Logic:** Significant parts of the task (accurate red flag detection via LLM, precise inline commenting with citations, checking for missing documents) are still implemented with placeholder logic and require further development.
*   **Checking for Missing Documents:** The specific requirement to check for mandatory *missing* documents based on the checklist is not yet implemented.
"""
